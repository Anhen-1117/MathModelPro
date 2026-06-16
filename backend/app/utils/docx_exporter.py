"""
Markdown → Word (.docx) 转换器

将数学建模论文的 Markdown 内容转换为 Word 文档，
保留标题层级、表格、代码块、列表、数学公式等结构。
支持图片嵌入和 LaTeX 数学公式转 Unicode。
"""

from __future__ import annotations

import re
import io
import os
import unicodedata
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── XML 安全工具 ──────────────────────────────────────────
_XML_ILLEGAL = re.compile(
    "[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x84\x86-\x9F"
    "\U0001FFFE-\U0001FFFF\U0002FFFE-\U0002FFFF"
    "\U0003FFFE-\U0003FFFF\U0004FFFE-\U0004FFFF"
    "\U0005FFFE-\U0005FFFF\U0006FFFE-\U0006FFFF"
    "\U0007FFFE-\U0007FFFF\U0008FFFE-\U0008FFFF"
    "\U0009FFFE-\U0009FFFF\U000AFFFE-\U000AFFFF"
    "\U000BFFFE-\U000BFFFF\U000CFFFE-\U000CFFFF"
    "\U000DFFFE-\U000DFFFF\U000EFFFE-\U000EFFFF"
    "\U000FFFFE-\U000FFFFF\U0010FFFE-\U0010FFFF]"
)

def _sanitize(text: str) -> str:
    """移除 XML 不兼容字符"""
    if not text:
        return text
    text = _XML_ILLEGAL.sub("", text)
    cleaned = []
    for ch in text:
        cat = unicodedata.category(ch)
        if cat == "Cc" and ch not in ("\n", "\r", "\t"):
            continue
        cleaned.append(ch)
    return "".join(cleaned)

def _safe_run(paragraph, text: str):
    return paragraph.add_run(_sanitize(text))

# ── 正则模式 ──────────────────────────────────────────────
RE_HEADING = re.compile(r"^(#{1,6})\s+(.+)$")
RE_UNORDERED_LIST = re.compile(r"^(\s*)[-*+]\s+(.+)$")
RE_ORDERED_LIST = re.compile(r"^(\s*)\d+\.\s+(.+)$")
RE_HORIZONTAL_RULE = re.compile(r"^(\s*)([-*_]){3,}\s*$")
RE_BLOCKQUOTE = re.compile(r"^>\s?(.*)$")
RE_TABLE_SEPARATOR = re.compile(r"^\s*\|?\s*[-:]+\s*(\|\s*[-:]+\s*)+\s*\|?\s*$")
RE_TABLE_ROW = re.compile(r"^\s*\|.+\|\s*$")
RE_CODE_FENCE_START = re.compile(r"^```(\w*)\s*$")
RE_CODE_FENCE_END = re.compile(r"^```\s*$")
RE_IMAGE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
RE_LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
RE_BOLD = re.compile(r"\*\*(.+?)\*\*")
RE_ITALIC = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")
RE_INLINE_CODE = re.compile(r"`([^`]+)`")
RE_INLINE_MATH = re.compile(r"(?<!\$)\$(?!\$)(.+?)(?<!\$)\$(?!\$)")

# ── 统一样式（从 paper_style 导入）──────────────────────
from app.utils.paper_style import get_style, FONTS as _F, SIZES as _S, PARAGRAPH as _P, HEADING as _H, TABLE as _T, CODE as _C

FONT_FAMILY = _F["body"]
FONT_FAMILY_ASIAN = _F["body_cjk"]
CODE_FONT = _F["code"][0]
MATH_FONT = _F["math"]

FONT_SIZES = {
    "h1": Pt(_S["h1"]), "h2": Pt(_S["h2"]), "h3": Pt(_S["h3"]),
    "h4": Pt(12), "h5": Pt(11), "h6": Pt(10.5),
    "body": Pt(_S["body"]), "code": Pt(_S["code"]),
    "math": Pt(_S["math"]), "table": Pt(_S["table"]),
}

def _set_font(run, name=FONT_FAMILY, size=None, bold=False, italic=False, color=None):
    run.font.name = name
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = size
    if color:
        run.font.color.rgb = color
    try:
        r = run._element
        rPr = r.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            r.insert(0, rPr)
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:eastAsia"), FONT_FAMILY_ASIAN)
    except Exception:
        pass

def _add_math_run(paragraph, text: str):
    run = _safe_run(paragraph, text)
    _set_font(run, MATH_FONT, FONT_SIZES["math"], italic=True)
    return run

def _add_code_run(paragraph, text: str):
    run = _safe_run(paragraph, text)
    _set_font(run, CODE_FONT, FONT_SIZES["code"], color=RGBColor(0x1A, 0x1A, 0x2E))
    return run

def _add_normal_run(paragraph, text: str):
    run = _safe_run(paragraph, text)
    _set_font(run, FONT_FAMILY, FONT_SIZES["body"])
    return run

# ── LaTeX 数学清洗 ─────────────────────────────────────────

# LaTeX 命令 → Unicode 映射表（按长度降序，长命令优先匹配）
_LATEX_SYMBOLS = [
    ("\\longrightarrow", "⟶"), ("\\longleftarrow", "⟵"),
    ("\\Longrightarrow", "⟹"), ("\\Longleftarrow", "⟸"),
    ("\\leftrightarrow", "↔"), ("\\Leftrightarrow", "⇔"),
    ("\\Rightarrow", "⇒"), ("\\Leftarrow", "⇐"),
    ("\\rightarrow", "→"), ("\\leftarrow", "←"),
    ("\\uparrow", "↑"), ("\\downarrow", "↓"),
    ("\\varepsilon", "ε"), ("\\varnothing", "∅"),
    ("\\boldsymbol", ""), ("\\mathbf", ""),
    ("\\mathrm", ""), ("\\mathit", ""),
    ("\\mathbb", ""), ("\\mathcal", ""),
    ("\\mathfrak", ""), ("\\mathsf", ""),
    ("\\mathtt", ""), ("\\mathscr", ""),
    ("\\partial", "∂"), ("\\infty", "∞"),
    ("\\nabla", "∇"), ("\\Delta", "Δ"),
    ("\\approx", "≈"), ("\\equiv", "≡"),
    ("\\neq", "≠"), ("\\propto", "∝"),
    ("\\leq", "≤"), ("\\geq", "≥"),
    ("\\ll", "≪"), ("\\gg", "≫"),
    ("\\sim", "∼"), ("\\simeq", "≃"),
    ("\\cdot", "·"), ("\\cdots", "⋯"),
    ("\\vdots", "⋮"), ("\\ddots", "⋱"),
    ("\\times", "×"), ("\\div", "÷"),
    ("\\pm", "±"), ("\\mp", "∓"),
    ("\\circ", "∘"), ("\\bullet", "•"),
    ("\\oplus", "⊕"), ("\\otimes", "⊗"),
    ("\\subset", "⊂"), ("\\supset", "⊃"),
    ("\\subseteq", "⊆"), ("\\supseteq", "⊇"),
    ("\\in", "∈"), ("\\notin", "∉"),
    ("\\forall", "∀"), ("\\exists", "∃"),
    ("\\emptyset", "∅"), ("\\varnothing", "∅"),
    ("\\cup", "∪"), ("\\cap", "∩"),
    ("\\setminus", "\\"), ("\\mid", "|"),
    ("\\alpha", "α"), ("\\beta", "β"),
    ("\\gamma", "γ"), ("\\delta", "δ"),
    ("\\epsilon", "ε"), ("\\zeta", "ζ"),
    ("\\eta", "η"), ("\\theta", "θ"),
    ("\\iota", "ι"), ("\\kappa", "κ"),
    ("\\lambda", "λ"), ("\\mu", "μ"),
    ("\\nu", "ν"), ("\\xi", "ξ"),
    ("\\pi", "π"), ("\\rho", "ρ"),
    ("\\sigma", "σ"), ("\\tau", "τ"),
    ("\\upsilon", "υ"), ("\\phi", "φ"),
    ("\\chi", "χ"), ("\\psi", "ψ"),
    ("\\omega", "ω"),
    ("\\Gamma", "Γ"), ("\\Theta", "Θ"),
    ("\\Lambda", "Λ"), ("\\Xi", "Ξ"),
    ("\\Pi", "Π"), ("\\Sigma", "Σ"),
    ("\\Upsilon", "Υ"), ("\\Phi", "Φ"),
    ("\\Psi", "Ψ"), ("\\Omega", "Ω"),
    ("\\sum", "Σ"), ("\\prod", "∏"),
    ("\\int", "∫"), ("\\oint", "∮"),
    ("\\hat{", ""), ("\\bar{", ""),
    ("\\tilde{", ""), ("\\vec{", ""),
    ("\\dot{", ""), ("\\ddot{", ""),
    ("\\sqrt", "√"), ("\\log", "log"),
    ("\\ln", "ln"), ("\\sin", "sin"),
    ("\\cos", "cos"), ("\\tan", "tan"),
    ("\\lim", "lim"), ("\\max", "max"),
    ("\\min", "min"), ("\\sup", "sup"),
    ("\\inf", "inf"), ("\\arg", "arg"),
    ("\\gcd", "gcd"), ("\\det", "det"),
    ("\\dim", "dim"), ("\\hom", "hom"),
    ("\\ker", "ker"), ("\\Pr", "Pr"),
    ("\\text", ""), ("\\frac", ""),
    ("\\left", ""), ("\\right", ""),
    ("\\ge", "≥"), ("\\le", "≤"),
    ("\\lVert", "‖"), ("\\rVert", "‖"),
    ("\\langle", "⟨"), ("\\rangle", "⟩"),
    ("\\lfloor", "⌊"), ("\\rfloor", "⌋"),
    ("\\lceil", "⌈"), ("\\rceil", "⌉"),
    ("\\quad", "  "), ("\\qquad", "    "),
    ("\\to", "→"), ("\\gets", "←"),
    ("\\mapsto", "↦"), ("\\iff", "⟺"),
    ("\\neg", "¬"), ("\\wedge", "∧"),
    ("\\vee", "∨"), ("\\vdash", "⊢"),
    ("\\|", "‖"),
    ("\\dashv", "⊣"), ("\\models", "⊨"),
    ("\\perp", "⊥"), ("\\parallel", "∥"),
    ("\\angle", "∠"), ("\\triangle", "△"),
    ("\\square", "□"), ("\\diamond", "◇"),
    ("\\star", "★"), ("\\ast", "∗"),
    ("\\ell", "ℓ"), ("\\hbar", "ħ"),
    ("\\imath", "ı"), ("\\jmath", "ȷ"),
    # 转义序列
    ("\\{", "{"), ("\\}", "}"),
    ("\\_", "_"), ("\\$", "$"),
    ("\\&", "&"), ("\\%", "%"),
    ("\\#", "#"), ("\\\\", "\n"),
]

def _parse_frac(s: str, start: int) -> tuple[str, int]:
    """解析 \\frac{num}{den} → (num)/(den)"""
    if not s.startswith("\\frac", start):
        return "", start
    i = start + 5  # 跳过 \frac
    # 跳过空白
    while i < len(s) and s[i] in " \t":
        i += 1
    if i >= len(s) or s[i] != "{":
        return "", start
    # 解析分子（匹配括号）
    depth = 1
    i += 1
    num_start = i
    while i < len(s) and depth > 0:
        if s[i] == "{":
            depth += 1
        elif s[i] == "}":
            depth -= 1
        i += 1
    num = s[num_start:i-1]
    # 跳过空白
    while i < len(s) and s[i] in " \t":
        i += 1
    if i >= len(s) or s[i] != "{":
        return "", start
    # 解析分母
    depth = 1
    i += 1
    den_start = i
    while i < len(s) and depth > 0:
        if s[i] == "{":
            depth += 1
        elif s[i] == "}":
            depth -= 1
        i += 1
    den = s[den_start:i-1]
    return f"({num})/({den})", i

def _replace_sqrt(text: str) -> str:
    """将 \\sqrt{...} 替换为 √(...)，正确匹配花括号嵌套"""
    result = []
    i = 0
    while i < len(text):
        if text[i:i+5] == "\\sqrt" and i+5 < len(text):
            j = i + 5
            # 跳过空白
            while j < len(text) and text[j] in " \t":
                j += 1
            if j < len(text) and text[j] == "{":
                # 匹配花括号
                depth = 1
                k = j + 1
                while k < len(text) and depth > 0:
                    if text[k] == "{":
                        depth += 1
                    elif text[k] == "}":
                        depth -= 1
                    k += 1
                inner = text[j+1:k-1]
                result.append(f"√({inner})")
                i = k
                continue
        result.append(text[i])
        i += 1
    return "".join(result)

def _clean_math(text: str) -> str:
    """将 LaTeX 数学公式转为可读的 Unicode 文本"""
    if not text:
        return text

    # 第1步：处理 \frac（需要在其他替换之前）
    result = []
    i = 0
    while i < len(text):
        frac_result, new_i = _parse_frac(text, i)
        if frac_result:
            result.append(frac_result)
            i = new_i
        else:
            result.append(text[i])
            i += 1
    text = "".join(result)

    # 第1.5步：\sqrt{...} → √(...)（保留括号表示范围）
    text = _replace_sqrt(text)

    # 第2步：替换 LaTeX 符号（按长度降序，确保 \int 先于 \in）
    sorted_symbols = sorted(_LATEX_SYMBOLS, key=lambda x: -len(x[0]))
    for cmd, replacement in sorted_symbols:
        text = text.replace(cmd, replacement)

    # 第3步：智能括号处理 — 保留上/下标结构，移除无意义的括号
    # 处理 _{...} → 保留下标内容，简化括号
    text = _unwrap_simple_braces(text)

    # 第4步：清理残留
    # 移除剩余的空花括号
    text = re.sub(r"\{\s*\}", "", text)
    # 移除孤立的 \left \right 残留
    text = text.replace("\\left", "").replace("\\right", "")
    # 修复词间粘连: argmin → arg min, max∑ → max ∑ 等
    text = re.sub(r"(arg|max|min|lim|log|ln|sin|cos|tan|sup|inf|det|gcd)([a-zA-Zα-ωΑ-Ω∑∏∫])", r"\1 \2", text)
    # 清理多余空白
    text = re.sub(r"\s+", " ", text).strip()

    return text

def _unwrap_simple_braces(text: str) -> str:
    """智能处理花括号：保留上/下标结构，移除无意义的单元素括号"""
    # 处理 _{...} 和 ^{...} 模式
    def _process_subscript(m):
        inner = m.group(1)
        # 如果是单个字符或数字，不需要括号
        if re.match(r"^[\w\dα-ωΑ-Ω∑∏∫∂∇∞±≤≥≠∈∀∃→←↑↓:.']$", inner):
            return f"_{inner}"
        # 多个元素，保留结构但去括号
        return f"_{inner}"

    def _process_superscript(m):
        inner = m.group(1)
        if re.match(r"^[\w\dα-ωΑ-Ω∑∏∫∂∇∞±≤≥≠∈∀∃→←↑↓:.']$", inner):
            return f"^{inner}"
        return f"^{inner}"

    text = re.sub(r"_{([^}]+)}", _process_subscript, text)
    text = re.sub(r"\^{([^}]+)}", _process_superscript, text)

    # 移除剩余的 { 和 }
    text = text.replace("{", "").replace("}", "")
    return text

# ── 格式化文本辅助（递归解析内联格式，叠加外层样式） ────

def _add_formatted_text(paragraph, text: str, bold=False, italic=False,
                        color=None, underline=False, project_dir=None):
    """将文本追加到段落，支持内联 $...$ 公式和 `...` 代码，并叠加外层样式"""
    if not text:
        return
    text = _sanitize(text)
    i = 0
    n = len(text)
    buf = ""

    def flush_buf():
        nonlocal buf
        if buf:
            run = _safe_run(paragraph, buf)
            _set_font(run, FONT_FAMILY, FONT_SIZES["body"], bold=bold, italic=italic, color=color)
            if underline:
                run.underline = True
            buf = ""

    while i < n:
        remaining = text[i:]

        # 内联代码
        m = RE_INLINE_CODE.match(remaining)
        if m:
            flush_buf()
            run = _safe_run(paragraph, m.group(1))
            _set_font(run, CODE_FONT, FONT_SIZES["code"], color=RGBColor(0x1A, 0x1A, 0x2E))
            run.bold = bold  # 保留外层粗体
            i += m.end()
            continue

        # 内联数学
        m = RE_INLINE_MATH.match(remaining)
        if m:
            flush_buf()
            math_text = _clean_math(m.group(1))
            if math_text:
                run = _safe_run(paragraph, math_text)
                _set_font(run, MATH_FONT, FONT_SIZES["math"], italic=True)
                if bold:
                    run.bold = True
            i += m.end()
            continue

        buf += text[i]
        i += 1

    flush_buf()


# ── 内联解析 ──────────────────────────────────────────────

def _parse_inline(paragraph, text: str, project_dir: str = None):
    """解析内联 Markdown 格式并追加到段落"""
    if not text:
        return
    text = _sanitize(text)
    i = 0
    n = len(text)
    buf = ""

    def flush_buf():
        nonlocal buf
        if buf:
            _add_formatted_text(paragraph, buf, project_dir=project_dir)
            buf = ""

    while i < n:
        remaining = text[i:]

        # 1. 图片 ![alt](url)
        m = RE_IMAGE.match(remaining)
        if m:
            flush_buf()
            alt_text = m.group(1)
            img_path = m.group(2)
            if project_dir and img_path:
                full_path = _resolve_image_path(img_path, project_dir)
                if full_path and os.path.isfile(full_path):
                    try:
                        run = paragraph.add_run()
                        run.add_picture(full_path, width=Inches(5.5))
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        i += m.end()
                        continue
                    except Exception:
                        pass
            run = _safe_run(paragraph, f"[图: {alt_text or img_path}]")
            _set_font(run, FONT_FAMILY, FONT_SIZES["body"], italic=True, color=RGBColor(0x88, 0x88, 0x88))
            i += m.end()
            continue

        # 2. 链接 [text](url)
        m = RE_LINK.match(remaining)
        if m:
            flush_buf()
            _add_formatted_text(paragraph, m.group(1), color=RGBColor(0x00, 0x66, 0xCC), underline=True, project_dir=project_dir)
            i += m.end()
            continue

        # 3. 粗体 **text**
        m = RE_BOLD.match(remaining)
        if m:
            flush_buf()
            _add_formatted_text(paragraph, m.group(1), bold=True, project_dir=project_dir)
            i += m.end()
            continue

        # 4. 斜体 *text*
        m = RE_ITALIC.match(remaining)
        if m:
            flush_buf()
            _add_formatted_text(paragraph, m.group(1), italic=True, project_dir=project_dir)
            i += m.end()
            continue

        # 5. 内联代码 `text`
        m = RE_INLINE_CODE.match(remaining)
        if m:
            flush_buf()
            _add_code_run(paragraph, m.group(1))
            i += m.end()
            continue

        # 6. 内联数学 $text$
        m = RE_INLINE_MATH.match(remaining)
        if m:
            flush_buf()
            math_text = _clean_math(m.group(1))
            if math_text:
                _add_math_run(paragraph, math_text)
            i += m.end()
            continue

        # 7. 普通字符
        buf += text[i]
        i += 1

    flush_buf()

def _resolve_image_path(img_path: str, project_dir: str) -> str | None:
    """解析图片路径，返回文件系统中的绝对路径"""
    # 去掉可能的 URL 前缀
    if img_path.startswith("http://") or img_path.startswith("https://"):
        # 尝试从路径中提取文件名
        filename = os.path.basename(img_path.split("?")[0])
    else:
        filename = img_path

    # 候选路径
    candidates = [
        os.path.join(project_dir, filename),
        os.path.join(project_dir, os.path.basename(filename)),
        os.path.join(project_dir, "figures", os.path.basename(filename)),
        # 直接使用 img_path（可能是绝对路径）
        img_path,
    ]
    for p in candidates:
        if os.path.isfile(p):
            return p
    return None

# ── 块级元素 ──────────────────────────────────────────────

def _add_paragraph(doc, text, style=None, project_dir=None):
    p = doc.add_paragraph(style=style)
    # 正文：首行缩进 2 字符 + 两端对齐（CUMCM 规范）
    if not style or style == "Normal":
        pf = p.paragraph_format
        pf.first_line_indent = Cm(0.74)  # 12pt 字号下约 2em
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if text:
        _parse_inline(p, text, project_dir)
    return p

def _apply_three_line_table(table):
    """将表格改为三线表样式（顶底粗线 + 表头细线）"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    # 移除默认边框
    for borders in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(borders)
    borders = OxmlElement("w:tblBorders")
    # 顶部粗线
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "12")  # 1.5pt
    top.set(qn("w:space"), "0")
    top.set(qn("w:color"), "000000")
    borders.append(top)
    # 底部粗线
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), "000000")
    borders.append(bottom)
    # 内部水平线（仅表头下方细线）
    insideH = OxmlElement("w:insideH")
    insideH.set(qn("w:val"), "single")
    insideH.set(qn("w:sz"), "6")  # 0.75pt
    insideH.set(qn("w:space"), "0")
    insideH.set(qn("w:color"), "000000")
    borders.append(insideH)
    tblPr.append(borders)
    # 移除单元格内部竖线：遍历所有单元格清除左右边框
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ["left", "right"]:
                el = OxmlElement(f"w:{side}")
                el.set(qn("w:val"), "nil")
                tcBorders.append(el)
            tcPr.append(tcBorders)


def _add_code_block(doc, lines):
    if not lines:
        return
    for line in lines:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Cm(0.7)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = Pt(15)
        # 背景色（浅灰）
        pPr = p._element.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "F2F2F2")
        shd.set(qn("w:val"), "clear")
        pPr.append(shd)
        run = _safe_run(p, line if line else " ")
        _set_font(run, CODE_FONT, FONT_SIZES["code"], color=RGBColor(0x1A, 0x1A, 0x2E))
    # 代码块后空行
    doc.add_paragraph()

def _add_table(doc, rows):
    if not rows:
        return
    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    # 三线表：顶底粗线 + 表头细线（CUMCM 规范）
    _apply_three_line_table(table)
    table.autofit = True
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j in range(num_cols):
            cell = row.cells[j]
            cell_text = row_data[j].strip() if j < len(row_data) else ""
            # 首段落：用 _parse_inline 处理公式/粗体/斜体
            first_p = cell.paragraphs[0]
            first_p.clear()
            _parse_inline(first_p, cell_text)
            # 设置单元格内所有段落间距
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(2)
                # 确保所有 run 字体正确
                for run in paragraph.runs:
                    if run.font.name is None:
                        _set_font(run, FONT_FAMILY, FONT_SIZES["table"])
    doc.add_paragraph()

def _parse_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [cell.strip() for cell in line.split("|")]

def _handle_standalone_image(line: str, doc, project_dir: str) -> bool:
    """处理独立行图片 ![alt](path)，返回 True 表示已处理"""
    m = RE_IMAGE.fullmatch(line.strip())
    if not m:
        return False
    alt = m.group(1)
    path = m.group(2)
    if project_dir:
        full = _resolve_image_path(path, project_dir)
        if full and os.path.isfile(full):
            try:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(full, width=Inches(5.5))
                # 图注
                if alt:
                    cap = doc.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap_run = _safe_run(cap, alt)
                    _set_font(cap_run, FONT_FAMILY, Pt(9), italic=True, color=RGBColor(0x66, 0x66, 0x66))
                return True
            except Exception:
                pass
    # 无法嵌入 → 保留文本说明
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = _safe_run(p, f"[图: {alt or path}]")
    _set_font(run, FONT_FAMILY, FONT_SIZES["body"], italic=True, color=RGBColor(0x88, 0x88, 0x88))
    return True

# ── 主转换函数 ────────────────────────────────────────────

def markdown_to_docx(markdown_text: str, title: str = "数学建模论文",
                     project_dir: str = None) -> bytes:
    """将 Markdown 论文转换为 .docx 字节流"""
    doc = Document()

    # ── 全局样式 ──────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = FONT_FAMILY
    style.font.size = FONT_SIZES["body"]
    style.paragraph_format.line_spacing = Pt(22)
    style.paragraph_format.space_after = Pt(6)
    try:
        style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_FAMILY_ASIAN)
    except Exception:
        pass

    for level in range(1, 5):
        name = f"Heading {level}"
        if name in doc.styles:
            hs = doc.styles[name]
            hs.font.name = FONT_FAMILY
            hs.font.size = FONT_SIZES.get(f"h{level}", Pt(12))
            hs.font.bold = True
            hs.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # ── 页面边距（统一 2.5cm） ──────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── 文档标题 ──────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(_safe_run(title_p, _sanitize(title)), FONT_FAMILY, Pt(22), bold=True)
    doc.add_paragraph()

    # ── 预处理与解析 ──────────────────────────────────────
    markdown_text = _sanitize(markdown_text)
    lines = markdown_text.split("\n")
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]

        # 空行
        if not line.strip():
            i += 1
            continue

        # 独立图片行 ![...](...)
        if _handle_standalone_image(line, doc, project_dir):
            i += 1
            continue

        # 代码块 ```
        m = RE_CODE_FENCE_START.match(line)
        if m:
            code_lines = []
            i += 1
            while i < n and not RE_CODE_FENCE_END.match(lines[i]):
                code_lines.append(lines[i])
                i += 1
            i += 1
            if code_lines:
                _add_code_block(doc, code_lines)
            continue

        # 数学块 $$
        if line.strip().startswith("$$"):
            rest = line.strip()[2:]
            end_idx = rest.find("$$")
            if end_idx != -1:
                formula = rest[:end_idx].strip()
                if formula:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _add_math_run(p, _clean_math(formula))
                i += 1
                continue
            else:
                math_lines = []
                i += 1
                while i < n and not lines[i].strip().startswith("$$"):
                    math_lines.append(lines[i])
                    i += 1
                i += 1
                if math_lines:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    formula = " ".join(ml.strip() for ml in math_lines if ml.strip())
                    cleaned = _clean_math(formula)
                    _add_math_run(p, cleaned)
                continue

        # 水平分割线
        if RE_HORIZONTAL_RULE.match(line):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            try:
                pPr = p._element.get_or_add_pPr()
                pBdr = OxmlElement("w:pBdr")
                bottom = OxmlElement("w:bottom")
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:sz"), "6")
                bottom.set(qn("w:space"), "4")
                bottom.set(qn("w:color"), "999999")
                pBdr.append(bottom)
                pPr.append(pBdr)
            except Exception:
                pass
            i += 1
            continue

        # 标题
        m = RE_HEADING.match(line)
        if m:
            level = min(len(m.group(1)), 4)
            heading_text = _sanitize(m.group(2).strip())
            if heading_text:
                p = doc.add_heading(heading_text, level=level)
                if level == 1:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    _set_font(run, FONT_FAMILY, FONT_SIZES.get(f"h{level}", Pt(12)),
                              bold=True, color=RGBColor(0x00, 0x00, 0x00))
            i += 1
            continue

        # 表格
        if RE_TABLE_ROW.match(line):
            table_rows = []
            while i < n:
                current = lines[i].strip()
                if RE_TABLE_ROW.match(current):
                    if not RE_TABLE_SEPARATOR.match(current):
                        table_rows.append(_parse_table_row(current))
                elif not current or RE_HEADING.match(current):
                    break
                else:
                    break
                i += 1
            if table_rows:
                _add_table(doc, table_rows)
            continue

        # 表格分隔行
        if RE_TABLE_SEPARATOR.match(line.strip()):
            i += 1
            continue

        # 无序列表
        m = RE_UNORDERED_LIST.match(line)
        if m:
            indent_level = len(m.group(1))
            text = m.group(2)
            p = doc.add_paragraph(style="List Bullet")
            p.clear()
            _parse_inline(p, text, project_dir)
            if indent_level > 0:
                p.paragraph_format.left_indent = Cm(1.27 * (1 + indent_level // 2))
            i += 1
            continue

        # 有序列表
        m = RE_ORDERED_LIST.match(line)
        if m:
            indent_level = len(m.group(1))
            text = m.group(2)
            p = doc.add_paragraph(style="List Number")
            p.clear()
            _parse_inline(p, text, project_dir)
            if indent_level > 0:
                p.paragraph_format.left_indent = Cm(1.27 * (1 + indent_level // 2))
            i += 1
            continue

        # 引用块
        m = RE_BLOCKQUOTE.match(line)
        if m:
            text = m.group(1)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.27)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            if text:
                _parse_inline(p, text, project_dir)
                for run in p.runs:
                    run.italic = True
                    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # 普通段落
        _add_paragraph(doc, line, project_dir=project_dir)
        i += 1

    # ── 保存 ──────────────────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
