"""文档解析工具 — 从 PDF / DOCX / DOC 提取文本（零成本，纯本地，稳定优先）"""

import os
import subprocess
import logging
from typing import Optional

logger = logging.getLogger(__name__)

# 支持解析为文本的扩展名
TEXT_EXTRACT_EXTENSIONS = {".pdf", ".docx", ".doc"}
# 数据文件扩展名（不需要提取文本）
DATA_FILE_EXTENSIONS = {".xlsx", ".csv", ".txt", ".xls"}


def extract_text_from_pdf(file_path: str) -> str:
    """使用 pdfplumber 提取 PDF 文本，逐页合并"""
    try:
        import pdfplumber
    except ImportError:
        logger.warning("pdfplumber 未安装，跳过 PDF 文本提取")
        return ""

    text_parts: list[str] = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return "\n\n".join(text_parts).strip()
    except Exception as e:
        logger.warning(f"PDF 文本提取失败 ({file_path}): {e}")
        return ""


def extract_text_from_docx(file_path: str) -> str:
    """使用 python-docx 提取 DOCX 段落文本"""
    try:
        from docx import Document
    except ImportError:
        logger.warning("python-docx 未安装，跳过 DOCX 文本提取")
        return ""

    try:
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # 也尝试提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells if cell.text.strip())
                if row_text.strip():
                    paragraphs.append(row_text)
        return "\n\n".join(paragraphs).strip()
    except Exception as e:
        logger.warning(f"DOCX 文本提取失败 ({file_path}): {e}")
        return ""


def extract_text_from_doc(file_path: str) -> str:
    """从旧格式 .doc 文件提取文本，多级降级策略"""
    # 策略 1: 尝试 antiword（Linux/Mac，最稳定）
    try:
        result = subprocess.run(
            ["antiword", file_path],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout.strip()
    except FileNotFoundError:
        pass  # antiword 未安装
    except Exception as e:
        logger.debug(f"antiword 提取失败: {e}")

    # 策略 2: 尝试用 python-docx 打开（某些 .doc 实际上是 .docx 格式）
    text = extract_text_from_docx(file_path)
    if text.strip():
        return text

    # 策略 3: 使用子进程调用 catdoc（另一个常见的 Linux .doc 工具）
    try:
        result = subprocess.run(
            ["catdoc", file_path],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout.strip()
    except FileNotFoundError:
        pass
    except Exception as e:
        logger.debug(f"catdoc 提取失败: {e}")

    # 策略 4: 从 OLE 二进制流中粗暴提取可读文本（最后降级）
    try:
        raw_text = _extract_raw_text_from_binary(file_path)
        if raw_text.strip():
            return raw_text
    except Exception as e:
        logger.debug(f"二进制原始提取失败: {e}")

    logger.warning(f"所有策略均无法提取 .doc 文本 ({file_path})，请考虑先转为 .docx")
    return ""


def _extract_raw_text_from_binary(file_path: str) -> str:
    """从二进制文件中提取可打印 ASCII/UTF-8 连续文本块（最低降级方案）"""
    try:
        with open(file_path, "rb") as f:
            data = f.read()
    except Exception:
        return ""

    # 尝试 UTF-8 解码，提取连续的可打印字符块
    result_parts: list[str] = []
    current_chunk: list[str] = []

    for byte in data:
        # 可打印 ASCII 范围 + 常见中文 UTF-8 多字节序列
        if 32 <= byte < 127 or byte in (10, 13):
            try:
                char = bytes([byte]).decode("ascii")
                current_chunk.append(char)
                continue
            except Exception:
                pass

    # 将收集的字符转为字符串
    raw = "".join(current_chunk)

    # 清理：移除过短的行和纯乱码行
    lines = raw.split("\n")
    clean_lines: list[str] = []
    for line in lines:
        stripped = line.strip()
        # 保留长度 > 3 且有字母或中文的行
        if len(stripped) > 3 and any(c.isalpha() or ord(c) > 127 for c in stripped):
            clean_lines.append(stripped)

    return "\n".join(clean_lines)


def parse_uploaded_file(file_path: str) -> dict:
    """
    解析上传的文件，返回提取结果。

    返回:
        {
            "filename": str,
            "extension": str,
            "extracted_text": str,    # 文本提取结果（仅 PDF/DOCX/DOC）
            "is_text_source": bool,   # 是否为题目文本来源
            "error": str | None       # 解析错误信息
        }
    """
    filename = os.path.basename(file_path)
    ext = os.path.splitext(filename)[1].lower()

    result = {
        "filename": filename,
        "extension": ext,
        "extracted_text": "",
        "is_text_source": ext in TEXT_EXTRACT_EXTENSIONS,
        "error": None,
    }

    if not os.path.isfile(file_path):
        result["error"] = f"文件不存在: {file_path}"
        return result

    if ext == ".pdf":
        result["extracted_text"] = extract_text_from_pdf(file_path)
    elif ext == ".docx":
        result["extracted_text"] = extract_text_from_docx(file_path)
    elif ext == ".doc":
        result["extracted_text"] = extract_text_from_doc(file_path)
    # xlsx/csv/txt 等数据文件不提取文本，保持原样

    if result["is_text_source"] and not result["extracted_text"]:
        result["error"] = "未能提取到文本内容（可能是扫描版 PDF 或旧格式文档）"

    return result


def merge_extracted_texts(parsed_files: list[dict]) -> str:
    """合并所有文件提取的文本为一份完整的题目描述"""
    texts: list[str] = []
    for pf in parsed_files:
        if pf["extracted_text"]:
            header = f"【{pf['filename']}】"
            texts.append(f"{header}\n{pf['extracted_text']}")
    return "\n\n---\n\n".join(texts)
